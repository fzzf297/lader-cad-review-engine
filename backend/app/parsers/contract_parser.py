"""
合同文件解析器 - 支持 Word 和 PDF 格式
"""
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


@dataclass
class ContractContent:
    """合同内容结构"""
    full_text: str = ""
    paragraphs: List[str] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class ContractParser:
    """合同文件解析器"""

    def parse(self, file_path: str) -> ContractContent:
        """根据文件类型选择解析器"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in ['.docx', '.doc']:
            return self._parse_word(file_path)
        elif suffix == '.pdf':
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    def _parse_word(self, file_path: str) -> ContractContent:
        """解析 Word 文档"""
        try:
            import docx
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = docx.Document(file_path)

        # 提取段落
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # 全文
        full_text = '\n'.join(paragraphs)

        return ContractContent(
            full_text=full_text,
            paragraphs=paragraphs,
            tables=tables,
            metadata={
                "file_type": "word",
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }
        )

    def _parse_pdf(self, file_path: str) -> ContractContent:
        """解析 PDF 文档"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")

        paragraphs = []
        tables = []
        page_count = 0

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                # 提取文本
                text = page.extract_text()
                if text:
                    paragraphs.extend([p.strip() for p in text.split('\n') if p.strip()])

                # 提取表格
                page_tables = page.extract_tables()
                tables.extend(page_tables or [])

        full_text = '\n'.join(paragraphs)

        return ContractContent(
            full_text=full_text,
            paragraphs=paragraphs,
            tables=tables,
            metadata={
                "file_type": "pdf",
                "page_count": page_count,
                "table_count": len(tables)
            }
        )


class WorkItemExtractor:
    """工作项提取器 - 使用关键词匹配提取合同工作项"""

    # 工作项关键词
    WORK_ITEM_KEYWORDS = {
        "门窗": ["门", "窗", "DOOR", "WINDOW", "铝合金", "塑钢", "木门", "钢门"],
        "管道": ["管", "PIPE", "给水", "排水", "暖通", "空调", "通风"],
        "电气": ["电", "开关", "插座", "灯具", "配电", "弱电", "ELEC"],
        "土建": ["混凝土", "钢筋", "砌体", "抹灰", "防水", "保温", "土方"],
        "装修": ["吊顶", "地面", "墙裙", "涂料", "油漆", "瓷砖", "地板"],
    }

    # 单位关键词
    UNIT_KEYWORDS = {
        "个": ["个", "套", "樘", "扇"],
        "米": ["米", "m", "延米", "延长米"],
        "平方米": ["平方米", "㎡", "平米", "平方"],
        "立方米": ["立方米", "m³", "立方"],
    }

    def extract_from_text(self, text: str) -> List[Dict[str, Any]]:
        """从文本中提取工作项（简化版，实际应使用 LLM）"""
        work_items = []

        # 遍历关键词，查找匹配
        for category, keywords in self.WORK_ITEM_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text:
                    # 简化的提取逻辑
                    work_items.append({
                        "name": keyword,
                        "category": category,
                        "quantity": 0,
                        "unit": "待确认",
                        "specification": "",
                        "source": "keyword_match",
                    })

        return work_items

    def extract_from_tables(self, tables: List[List[List[str]]]) -> List[Dict[str, Any]]:
        """从表格中提取工作项"""
        work_items = []

        for table in tables:
            if not table or len(table) < 2:
                continue

            # 查找表头
            header = table[0]
            header_lower = [h.lower() if h else "" for h in header]

            # 常见表头关键词
            name_idx = self._find_column(header_lower, ["名称", "项目", "工作内容", "item", "name"])
            qty_idx = self._find_column(header_lower, ["数量", "工程量", "qty", "quantity"])
            unit_idx = self._find_column(header_lower, ["单位", "unit"])
            spec_idx = self._find_column(header_lower, ["规格", "型号", "spec", "specification"])

            if name_idx is not None:
                for row in table[1:]:
                    if len(row) > name_idx and row[name_idx]:
                        item = {
                            "name": row[name_idx] if name_idx < len(row) else "",
                            "category": self._categorize(row[name_idx] if name_idx < len(row) else ""),
                            "quantity": self._parse_number(row[qty_idx] if qty_idx is not None and qty_idx < len(row) else "0"),
                            "unit": row[unit_idx] if unit_idx is not None and unit_idx < len(row) else "",
                            "specification": row[spec_idx] if spec_idx is not None and spec_idx < len(row) else "",
                            "source": "table_extract",
                        }
                        if item["name"]:
                            work_items.append(item)

        return work_items

    def _find_column(self, header: List[str], keywords: List[str]) -> Optional[int]:
        """查找包含关键词的列索引"""
        for i, h in enumerate(header):
            for kw in keywords:
                if kw in h:
                    return i
        return None

    def _categorize(self, name: str) -> str:
        """根据名称判断分类"""
        name_upper = name.upper()
        for category, keywords in self.WORK_ITEM_KEYWORDS.items():
            for kw in keywords:
                if kw.upper() in name_upper:
                    return category
        return "其他"

    def _parse_number(self, text: str) -> float:
        """解析数字"""
        try:
            # 移除可能的非数字字符
            text = text.replace(",", "").replace("，", "").strip()
            return float(text)
        except (ValueError, AttributeError):
            return 0